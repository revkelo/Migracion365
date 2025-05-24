# ---------------------------------------------------------------
# Script Fusionado: Exportación y Descarga Estructurada de Google Drive
# Autor: Kevin Gonzalez
# Descripción:
#   Este script descarga todos los archivos de una cuenta de Google Drive,
#   respetando la estructura de carpetas original. Exporta los archivos de
#   Google Docs, Sheets, Slides y Forms a sus formatos compatibles (.docx,
#   .xlsx, .pptx), mientras que los archivos binarios se descargan tal cual.
# ---------------------------------------------------------------

import os, io, re, pickle, time
from tqdm import tqdm
from docx import Document
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

def obtener_credenciales():
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as f:
            cred = pickle.load(f)
    else:
        cred = None
    if not cred or not cred.valid:
        if cred and cred.expired and cred.refresh_token:
            cred.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            cred = flow.run_local_server(port=8089)
        with open('token.pickle', 'wb') as f:
            pickle.dump(cred, f)
    return cred

def sanitizar(nombre):
    return re.sub(r'[\\/*?:"<>|]', "_", nombre)

def obtener_servicios():
    creds = obtener_credenciales()
    return build('drive', 'v3', credentials=creds), build('forms', 'v1', credentials=creds)

def crear_word(form, ruta):
    doc = Document()
    doc.add_heading(form.get('info', {}).get('title', 'Sin título'), 0)
    letras = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    for i, item in enumerate(form.get('items', []), 1):
        doc.add_paragraph(f"{i}. {item.get('title', 'Sin título')}")
        pregunta = item.get('questionItem', {}).get('question', {})
        if 'choiceQuestion' in pregunta:
            for j, op in enumerate(pregunta['choiceQuestion'].get('options', [])):
                doc.add_paragraph(f"{letras[j]}. {op.get('value', 'Opción')}")
        elif 'textQuestion' in pregunta:
            doc.add_paragraph("\n\n")
    doc.save(ruta)
    tqdm.write(f"Generado: {ruta}")

EXPORT_MIME = {
    'application/vnd.google-apps.document': ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx'),
    'application/vnd.google-apps.spreadsheet': ('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'xlsx'),
    'application/vnd.google-apps.presentation': ('application/vnd.openxmlformats-officedocument.presentationml.presentation', 'pptx'),
    'application/vnd.google-apps.form': ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx')  # Especial
}

def descargar_archivo(file, path, drive_service, forms_service):
    nombre = sanitizar(file['name'])
    ruta_local = os.path.join(path, nombre)

    try:
        if file['mimeType'] in EXPORT_MIME and file['mimeType'] != 'application/vnd.google-apps.form':
            export_mime, ext = EXPORT_MIME[file['mimeType']]
            request = drive_service.files().export_media(fileId=file['id'], mimeType=export_mime)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            while not downloader.next_chunk()[1]: pass
            with open(f"{ruta_local}.{ext}", 'wb') as f:
                f.write(fh.getvalue())
        elif file['mimeType'] == 'application/vnd.google-apps.form':
            detalle = forms_service.forms().get(formId=file['id']).execute()
            crear_word(detalle, f"{ruta_local}.docx")
        elif not file['mimeType'].startswith('application/vnd.google-apps.folder'):
            request = drive_service.files().get_media(fileId=file['id'])
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            while not downloader.next_chunk()[1]: pass
            with open(ruta_local, 'wb') as f:
                f.write(fh.getvalue())
    except Exception as e:
        tqdm.write(f"Error al descargar {nombre}: {e}")

def construir_arbol(drive_service):
    carpetas = {}
    archivos = {}
    page_token = None

    while True:
        res = drive_service.files().list(q="trashed = false", fields="nextPageToken, files(id, name, mimeType, parents)", pageToken=page_token).execute()
        for f in res['files']:
            if f['mimeType'] == 'application/vnd.google-apps.folder':
                carpetas[f['id']] = f
            else:
                archivos[f['id']] = f
        page_token = res.get('nextPageToken', None)
        if not page_token:
            break
    return carpetas, archivos

def obtener_ruta(file_id, carpetas):
    ruta = []
    actual = file_id
    while actual in carpetas:
        carpeta = carpetas[actual]
        ruta.append(sanitizar(carpeta['name']))
        padres = carpeta.get('parents')
        if not padres:
            break
        actual = padres[0]
    return list(reversed(ruta))

def ejecutar_descarga():
    drive, forms = obtener_servicios()
    carpetas, archivos = construir_arbol(drive)
    descarga_base = 'Descarga_GoogleDrive'
    os.makedirs(descarga_base, exist_ok=True)

    for archivo_id, archivo in tqdm(archivos.items(), desc="Procesando archivos", unit="file"):
        padres = archivo.get('parents')
        if padres:
            ruta = obtener_ruta(padres[0], carpetas)
        else:
            ruta = []
        path = os.path.join(descarga_base, *ruta)
        os.makedirs(path, exist_ok=True)
        descargar_archivo(archivo, path, drive, forms)

    print("Descarga completada.")

if __name__ == "__main__":
    ejecutar_descarga()
