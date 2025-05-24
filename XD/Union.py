#!/usr/bin/env python3
# ---------------------------------------------------------------
# Script de Migración Directa: Google Drive → OneDrive
# Autor: Kevin Gonzalez (Enhanced)
# Descripción:
#   Migra archivos directamente de Google Drive a OneDrive sin
#   almacenamiento local intermedio. Mantiene la estructura de
#   carpetas y maneja archivos de Google Workspace.
# ---------------------------------------------------------------

import os
import io
import re
import json
import time
import pickle
import logging
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from tqdm import tqdm

# Google Drive imports
from docx import Document
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError

# OneDrive imports
import msal
import requests

# ============================================================================
# CONFIGURACIÓN
# ============================================================================

# Google Drive
GOOGLE_SCOPES = [
    'https://www.googleapis.com/auth/drive.readonly',
    'https://www.googleapis.com/auth/forms.body.readonly'
]


# OneDrive
ONEDRIVE_CLIENT_ID = "d8227c52-3dde-4198-81a8-60f1347357ab"
ONEDRIVE_AUTHORITY = "https://login.microsoftonline.com/common"
ONEDRIVE_SCOPES = ["Files.ReadWrite.All"]

# Configuración general
CHUNK_SIZE = 10 * 1024 * 1024  # 10MB para uploads grandes
LARGE_FILE_THRESHOLD = 4 * 1024 * 1024  # 4MB umbral para chunked upload
MAX_RETRIES = 3
RETRY_DELAY = 2
LOG_FILE = 'migration_log.txt'
PROGRESS_FILE = 'migration_progress.json'

# Formatos de exportación de Google Workspace
GOOGLE_EXPORT_FORMATS = {
    'application/vnd.google-apps.document': {
        'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'ext': 'docx'
    },
    'application/vnd.google-apps.spreadsheet': {
        'mime': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'ext': 'xlsx'
    },
    'application/vnd.google-apps.presentation': {
        'mime': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'ext': 'pptx'
    },
    'application/vnd.google-apps.form': {
        'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'ext': 'docx'
    }
}

class DirectMigrator:
    def __init__(self, onedrive_folder: str = ""):
        self.onedrive_folder = onedrive_folder
        
        # Servicios
        self.google_drive = None
        self.google_forms = None
        self.onedrive_token = None
        
        # Setup logging
        self.setup_logging()
        
        # Progreso
        self.progress = self.load_progress()
        self.stats = {
            'total_files': 0,
            'migrated': 0,
            'skipped': 0,
            'errors': 0,
            'total_size': 0,
            'migrated_size': 0,
            'start_time': datetime.now()
        }
    
    def setup_logging(self):
        """Configurar logging"""
        log_format = '%(asctime)s - %(levelname)s - %(message)s'
        
        file_handler = logging.FileHandler(LOG_FILE, encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        file_handler.setFormatter(logging.Formatter(log_format))
        
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
        
        self.logger = logging.getLogger('DirectMigrator')
        self.logger.setLevel(logging.INFO)
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
    
    def load_progress(self) -> Dict:
        """Cargar progreso previo"""
        if Path(PROGRESS_FILE).exists():
            try:
                with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if 'migrated_files' in data and isinstance(data['migrated_files'], list):
                        data['migrated_files'] = set(data['migrated_files'])
                    return data
            except Exception as e:
                self.logger.warning(f"No se pudo cargar el progreso: {e}")
        return {'migrated_files': set()}
    
    def save_progress(self):
        """Guardar progreso actual"""
        try:
            progress_data = {
                'migrated_files': list(self.progress['migrated_files']),
                'last_updated': datetime.now().isoformat(),
                'stats': self.stats.copy()
            }
            if 'start_time' in progress_data['stats']:
                progress_data['stats']['start_time'] = progress_data['stats']['start_time'].isoformat()
            
            with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
                json.dump(progress_data, f, indent=2)
        except Exception as e:
            self.logger.error(f"No se pudo guardar el progreso: {e}")
    
    # ========================================================================
    # AUTENTICACIÓN GOOGLE DRIVE
    # ========================================================================
    
    def setup_google_services(self):
        """Configurar servicios de Google Drive"""
        creds = None
        token_file = 'token.pickle'
        
        if os.path.exists(token_file):
            os.remove(token_file)
            self.logger.info("🔁 Eliminado token.pickle para forzar reautenticación")

        creds = None  
        
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                if not Path('credentials.json').exists():
                    raise FileNotFoundError("credentials.json no encontrado")
                
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', GOOGLE_SCOPES)
                creds = flow.run_local_server(port=8089)
            
            with open(token_file, 'wb') as f:
                pickle.dump(creds, f)
        
        self.google_drive = build('drive', 'v3', credentials=creds)
        self.google_forms = build('forms', 'v1', credentials=creds)
        self.logger.info("Servicios de Google Drive configurados")
    
    # ========================================================================
    # AUTENTICACIÓN ONEDRIVE
    # ========================================================================
    
    def setup_onedrive_token(self):
        """Obtener token de OneDrive"""
        self.logger.info("Obteniendo token de OneDrive...")
        
        app = msal.PublicClientApplication(ONEDRIVE_CLIENT_ID, authority=ONEDRIVE_AUTHORITY)
        
        # Limpiar cuentas cacheadas
        for account in app.get_accounts():
            app.remove_account(account)
        
        result = app.acquire_token_interactive(ONEDRIVE_SCOPES, prompt="select_account")
        
        if "access_token" in result:
            self.onedrive_token = result["access_token"]
            self.logger.info("Token de OneDrive obtenido")
        else:
            raise Exception(f"Error obteniendo token: {result.get('error_description')}")
    
    # ========================================================================
    # FUNCIONES DE UTILIDAD
    # ========================================================================
    
    def sanitize_filename(self, filename: str) -> str:
        """Limpiar nombre de archivo"""
        sanitized = re.sub(r'[\\/*?:"<>|]', '_', filename)
        if len(sanitized) > 250:
            name, ext = os.path.splitext(sanitized)
            sanitized = name[:245] + ext
        return sanitized.strip()
    
    def format_size(self, size_bytes: int) -> str:
        """Formatear tamaño de archivo"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    # ========================================================================
    # MANEJO DE GOOGLE DRIVE
    # ========================================================================
    
    def build_google_structure(self) -> Tuple[Dict, Dict]:
        """Construir estructura de carpetas y archivos de Google Drive"""
        self.logger.info("Analizando estructura de Google Drive...")
        
        folders = {}
        files = {}
        page_token = None
        
        while True:
            try:
                result = self.google_drive.files().list(
                    q="trashed = false",
                    fields="nextPageToken, files(id, name, mimeType, parents, size, modifiedTime)",
                    pageSize=1000,
                    pageToken=page_token
                ).execute()
                
                for file_item in result['files']:
                    if file_item['mimeType'] == 'application/vnd.google-apps.folder':
                        folders[file_item['id']] = file_item
                    else:
                        files[file_item['id']] = file_item
                        # Calcular tamaño total
                        if 'size' in file_item:
                            self.stats['total_size'] += int(file_item['size'])
                
                page_token = result.get('nextPageToken')
                if not page_token:
                    break
                    
            except Exception as e:
                self.logger.error(f"Error construyendo estructura: {e}")
                raise
        
        self.stats['total_files'] = len(files)
        self.logger.info(f"Encontrados {len(folders)} carpetas y {len(files)} archivos")
        return folders, files
    
    def get_folder_path(self, parent_id: str, folders: Dict) -> List[str]:
        """Obtener ruta completa de carpeta"""
        path = []
        current_id = parent_id
        
        while current_id in folders:
            folder = folders[current_id]
            path.append(self.sanitize_filename(folder['name']))
            parents = folder.get('parents')
            if not parents:
                break
            current_id = parents[0]
        
        return list(reversed(path))
    
    def download_google_file_to_memory(self, file_info: Dict) -> Tuple[Optional[io.BytesIO], str]:
        """Descargar archivo de Google Drive a memoria"""
        file_id = file_info['id']
        mime_type = file_info['mimeType']
        file_name = file_info['name']
        
        try:
            # Manejar archivos de Google Workspace
            if mime_type in GOOGLE_EXPORT_FORMATS:
                if mime_type == 'application/vnd.google-apps.form':
                    # Crear documento Word desde Google Form
                    form_data = self.google_forms.forms().get(formId=file_id).execute()
                    return self.create_word_from_form(form_data), f"{self.sanitize_filename(file_name)}_forms.docx"

                else:
                    # Exportar otros tipos de Google Workspace
                    export_info = GOOGLE_EXPORT_FORMATS[mime_type]
                    request = self.google_drive.files().export_media(
                        fileId=file_id, 
                        mimeType=export_info['mime']
                    )
                    
                    fh = io.BytesIO()
                    downloader = MediaIoBaseDownload(fh, request)
                    
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                    
                    fh.seek(0)
                    return fh, f"{self.sanitize_filename(file_name)}.{export_info['ext']}"
            
            # Archivos binarios normales
            else:
                request = self.google_drive.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                
                fh.seek(0)
                return fh, self.sanitize_filename(file_name)
                
        except HttpError as e:
            if e.resp.status == 403:
                self.logger.warning(f"Sin permisos para descargar: {file_name}")
            elif e.resp.status == 404:
                self.logger.warning(f"Archivo no encontrado: {file_name}")
            else:
                self.logger.error(f"Error HTTP {e.resp.status} descargando {file_name}: {e}")
            return None, file_name
            
        except Exception as e:
            self.logger.error(f"Error descargando {file_name}: {e}")
            return None, file_name
    
    def create_word_from_form(self, form_data: Dict) -> io.BytesIO:
        """Crear documento Word desde datos de Google Form"""
        doc = Document()
        
        title = form_data.get('info', {}).get('title', 'Formulario sin título')
        doc.add_heading(title, 0)
        
        description = form_data.get('info', {}).get('description', '')
        if description:
            doc.add_paragraph(description)
            doc.add_paragraph()
        
        letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        
        for i, item in enumerate(form_data.get('items', []), 1):
            question_title = item.get('title', 'Pregunta sin título')
            doc.add_paragraph(f"{i}. {question_title}", style='Heading 2')
            
            question = item.get('questionItem', {}).get('question', {})
            
            if 'choiceQuestion' in question:
                options = question['choiceQuestion'].get('options', [])
                for j, option in enumerate(options):
                    prefix = f"{letters[j]}." if j < len(letters) else f"{j+1}."
                    doc.add_paragraph(f"  {prefix} {option.get('value', 'Opción')}")
            
            elif 'textQuestion' in question:
                paragraph_style = question['textQuestion'].get('paragraph', False)
                if paragraph_style:
                    doc.add_paragraph("  [Respuesta de texto largo]")
                else:
                    doc.add_paragraph("  [Respuesta de texto corto]")
            
            doc.add_paragraph()
        
        # Guardar en BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # ========================================================================
    # MANEJO DE ONEDRIVE
    # ========================================================================
    
    def create_onedrive_folder(self, folder_path: str) -> bool:
        """Crear estructura de carpetas en OneDrive"""
        if not folder_path.strip():
            return True
        
        headers = {"Authorization": f"Bearer {self.onedrive_token}"}
        path_parts = folder_path.strip('/').split('/')
        current_path = ""
        
        for part in path_parts:
            if not part:
                continue
            
            current_path = f"{current_path}/{part}" if current_path else part
            
            # Verificar si existe
            check_url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{(self.onedrive_folder + '/' + current_path).strip('/')}"
            response = requests.get(check_url, headers=headers)
            
            if response.status_code == 404:
                # Crear carpeta
                parent_path = f"{self.onedrive_folder}/{'/'.join(path_parts[:path_parts.index(part)])}" if path_parts.index(part) > 0 else self.onedrive_folder
                parent_path = parent_path.strip('/')
                
                create_url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{parent_path}:/children"
                folder_data = {
                    "name": part,
                    "folder": {},
                    "@microsoft.graph.conflictBehavior": "rename"
                }
                
                create_response = requests.post(create_url, headers=headers, json=folder_data)
                if create_response.status_code not in [200, 201]:
                    self.logger.error(f"Error creando carpeta {part}: {create_response.text}")
                    return False
        
        return True
    
    def upload_to_onedrive(self, file_data: io.BytesIO, remote_path: str, file_size: int) -> bool:
        """Subir archivo a OneDrive"""
        headers = {"Authorization": f"Bearer {self.onedrive_token}"}
        
        try:
            if file_size > LARGE_FILE_THRESHOLD:
                return self.upload_large_to_onedrive(file_data, remote_path, file_size, headers)
            else:
                return self.upload_small_to_onedrive(file_data, remote_path, headers)
                
        except Exception as e:
            self.logger.error(f"Error subiendo a OneDrive: {e}")
            return False
    
    def upload_small_to_onedrive(self, file_data: io.BytesIO, remote_path: str, headers: Dict) -> bool:
        """Subir archivo pequeño a OneDrive"""
        upload_headers = headers.copy()
        upload_headers["Content-Type"] = "application/octet-stream"
        
        url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{remote_path}:/content"
        response = requests.put(url, headers=upload_headers, data=file_data.read())
        
        return response.status_code in [200, 201]
    
    def upload_large_to_onedrive(self, file_data: io.BytesIO, remote_path: str, file_size: int, headers: Dict) -> bool:
        """Subir archivo grande a OneDrive con chunked upload"""
        # Crear sesión de upload
        session_url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{remote_path}:/createUploadSession"
        session_data = {"item": {"@microsoft.graph.conflictBehavior": "replace"}}
        
        session_response = requests.post(session_url, headers=headers, json=session_data)
        if session_response.status_code != 200:
            self.logger.error(f"Error creando sesión de upload: {session_response.text}")
            return False
        
        upload_url = session_response.json()["uploadUrl"]
        
        # Subir en chunks
        chunk_start = 0
        file_data.seek(0)
        
        while chunk_start < file_size:
            chunk_end = min(chunk_start + CHUNK_SIZE - 1, file_size - 1)
            chunk_data = file_data.read(chunk_end - chunk_start + 1)
            
            chunk_headers = {
                "Content-Length": str(len(chunk_data)),
                "Content-Range": f"bytes {chunk_start}-{chunk_end}/{file_size}"
            }
            
            chunk_response = requests.put(upload_url, headers=chunk_headers, data=chunk_data)
            
            if chunk_response.status_code not in [202, 200, 201]:
                self.logger.error(f"Error subiendo chunk: {chunk_response.text}")
                return False
            
            chunk_start = chunk_end + 1
        
        return True
    
    # ========================================================================
    # PROCESO PRINCIPAL DE MIGRACIÓN
    # ========================================================================
    
    def migrate_file_direct(self, file_info: Dict, remote_path: str) -> bool:
        """Migrar archivo directamente de Google Drive a OneDrive"""
        file_name = file_info['name']
        
        for attempt in range(MAX_RETRIES):
            try:
                # Descargar de Google Drive a memoria
                file_data, final_name = self.download_google_file_to_memory(file_info)
                
                if file_data is None:
                    return False
                
                # Calcular tamaño
                file_data.seek(0, 2)  # Ir al final
                file_size = file_data.tell()
                file_data.seek(0)  # Volver al inicio
                
                # Construir ruta remota final
                base_path = self.onedrive_folder.strip('/')
                final_remote_path = f"{base_path}/{remote_path}/{final_name}".strip('/').replace('//', '/')
                
                # Subir a OneDrive
                if self.upload_to_onedrive(file_data, final_remote_path, file_size):
                    self.stats['migrated_size'] += file_size
                    return True
                elif attempt < MAX_RETRIES - 1:
                    self.logger.warning(f"Reintentando migración de {file_name}...")
                    time.sleep(RETRY_DELAY * (attempt + 1))
                
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    self.logger.warning(f"Error migrando {file_name}, reintentando: {e}")
                    time.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    self.logger.error(f"Falló migración de {file_name} después de {MAX_RETRIES} intentos: {e}")
        
        return False
    
    def run_migration(self, skip_existing: bool = True):
        """Ejecutar migración completa"""
        try:
            self.logger.info("Iniciando migración directa Google Drive → OneDrive")
            
            # Configurar servicios
            print("🔐 Configurando Google Drive...")
            self.setup_google_services()
            
            print("🔐 Configurando OneDrive...")
            self.setup_onedrive_token()
            
            # Analizar estructura
            print("📁 Analizando estructura de Google Drive...")
            folders, files = self.build_google_structure()
            
            if not files:
                print("No se encontraron archivos para migrar.")
                return
            
            print(f"📊 Encontrados {len(files)} archivos ({self.format_size(self.stats['total_size'])})")
            
            # Migrar archivos
            with tqdm(total=len(files), desc="Migrando archivos", unit="archivo") as pbar:
                for file_id, file_info in files.items():
                    file_name = file_info['name']
                    mime_type = file_info['mimeType']
                    
                    # Saltar si ya fue migrado
                    if skip_existing and file_id in self.progress['migrated_files']:
                        self.stats['skipped'] += 1
                        pbar.update(1)
                        continue
                    
                    # Saltar tipos no soportados
                    if mime_type in ['application/vnd.google-apps.script', 'application/vnd.google-apps.site']:
                        self.logger.info(f"Saltando tipo no soportado: {file_name}")
                        self.progress['migrated_files'].add(file_id)
                        self.stats['skipped'] += 1
                        pbar.update(1)
                        continue
                    
                    # Construir ruta
                    parents = file_info.get('parents')
                    if parents:
                        folder_path = '/'.join(self.get_folder_path(parents[0], folders))
                    else:
                        folder_path = ""
                    
                    # Crear estructura de carpetas en OneDrive
                    if folder_path and not self.create_onedrive_folder(folder_path):
                        self.logger.error(f"Error creando estructura para {file_name}")
                        self.stats['errors'] += 1
                        pbar.update(1)
                        continue
                    
                    # Actualizar barra de progreso
                    pbar.set_description(f"Migrando: {file_name[:30]}...")
                    
                    # Migrar archivo
                    if self.migrate_file_direct(file_info, folder_path):
                        self.stats['migrated'] += 1
                        self.progress['migrated_files'].add(file_id)
                        self.logger.info(f"✅ Migrado: {folder_path}/{file_name}")
                    else:
                        self.stats['errors'] += 1
                        self.logger.error(f"❌ Error migrando: {file_name}")
                    
                    pbar.update(1)
                    
                    # Guardar progreso periódicamente
                    if (self.stats['migrated'] + self.stats['errors']) % 5 == 0:
                        self.save_progress()
            
            # Guardar progreso final
            self.save_progress()
            self.print_statistics()
            
        except KeyboardInterrupt:
            self.logger.info("Migración interrumpida por el usuario")
            self.save_progress()
            print("\n⏸️  Migración interrumpida. El progreso se ha guardado.")
            
        except Exception as e:
            self.logger.error(f"Error en la migración: {e}")
            self.save_progress()
            raise
    
    def print_statistics(self):
        """Mostrar estadísticas de migración"""
        duration = datetime.now() - self.stats['start_time']
        
        print("\n" + "="*60)
        print("ESTADÍSTICAS DE MIGRACIÓN")
        print("="*60)
        print(f"Total de archivos: {self.stats['total_files']}")
        print(f"Migrados exitosamente: {self.stats['migrated']}")
        print(f"Saltados: {self.stats['skipped']}")
        print(f"Errores: {self.stats['errors']}")
        print(f"Tamaño total: {self.format_size(self.stats['total_size'])}")
        print(f"Tamaño migrado: {self.format_size(self.stats['migrated_size'])}")
        print(f"Duración: {duration}")
        print(f"Carpeta destino: {self.onedrive_folder}")
        print("="*60)

def main():
    """Función principal"""
    print("🚀 Migración Directa Google Drive → OneDrive")
    print("="*50)
    print("📁 Carpeta destino: GoogleDrive_Backup")
    print("⏭️  Saltando archivos ya migrados automáticamente")
    print("="*50)
    
    # Ejecutar migración con configuración fija
    try:
        migrator = DirectMigrator("")
        migrator.run_migration(skip_existing=True)
        print("\n🎉 ¡Migración completada!")
        
    except Exception as e:
        print(f"\n❌ Error en la migración: {e}")

if __name__ == "__main__":
    main()