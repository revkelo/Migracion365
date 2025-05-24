import os
import io
import pickle
import logging
from pathlib import Path
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from docx import Document
from config import GOOGLE_SCOPES, GOOGLE_EXPORT_FORMATS
from utils import sanitize_filename

class GoogleService:
    def __init__(self, credentials_path: str = 'credentials.json', token_path: str = 'token.pickle'):
        self.credentials_path = credentials_path
        self.token_path = token_path
        self.drive = None
        self.forms = None
        self.logger = logging.getLogger('GoogleService')
        self.setup_services()

    def setup_services(self):
        creds = None
        if os.path.exists(self.token_path):
            os.remove(self.token_path)
            self.logger.info("🔁 Eliminado token.pickle para forzar reautenticación")

        if not creds or not getattr(creds, 'valid', False):
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                if not Path(self.credentials_path).exists():
                    raise FileNotFoundError(f"{self.credentials_path} no encontrado")
                flow = InstalledAppFlow.from_client_secrets_file(self.credentials_path, GOOGLE_SCOPES)
                creds = flow.run_local_server(port=8089)
            with open(self.token_path, 'wb') as f:
                pickle.dump(creds, f)

        self.drive = build('drive', 'v3', credentials=creds)
        self.forms = build('forms', 'v1', credentials=creds)
        self.logger.info("Servicios de Google Drive configurados")

    def list_files_and_folders(self):
        folders, files = {}, {}
        page_token = None
        total_size = 0
        while True:
            result = self.drive.files().list(
                q="trashed = false",
                fields="nextPageToken, files(id, name, mimeType, parents, size, modifiedTime)",
                pageSize=1000,
                pageToken=page_token
            ).execute()
            for item in result.get('files', []):
                if item['mimeType'] == 'application/vnd.google-apps.folder':
                    folders[item['id']] = item
                else:
                    files[item['id']] = item
                    total_size += int(item.get('size', 0))
            page_token = result.get('nextPageToken')
            if not page_token:
                break
        return folders, files, total_size

    def get_folder_path(self, parent_id: str, folders: dict) -> list:
        path, current = [], parent_id
        while current in folders:
            folder = folders[current]
            path.append(sanitize_filename(folder['name']))
            parents = folder.get('parents')
            current = parents[0] if parents else None
        return list(reversed(path))

    def download_file(self, file_info: dict):
        file_id, mime, name = file_info['id'], file_info['mimeType'], file_info['name']
        try:
            if mime in GOOGLE_EXPORT_FORMATS:
                exp = GOOGLE_EXPORT_FORMATS[mime]
                if mime == 'application/vnd.google-apps.form':
                    form_data = self.forms.forms().get(formId=file_id).execute()
                    return self.create_word_from_form(form_data), f"{sanitize_filename(name)}_forms.docx"
                req = self.drive.files().export_media(fileId=file_id, mimeType=exp['mime'])
                fh = io.BytesIO()
                dl = MediaIoBaseDownload(fh, req)
                done = False
                while not done:
                    _, done = dl.next_chunk()
                fh.seek(0)
                return fh, f"{sanitize_filename(name)}.{exp['ext']}"
            req = self.drive.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            dl = MediaIoBaseDownload(fh, req)
            done = False
            while not done:
                _, done = dl.next_chunk()
            fh.seek(0)
            return fh, sanitize_filename(name)
        except Exception as e:
            self.logger.error(f"Error descargando {name}: {e}")
            return None, name

    def create_word_from_form(self, form_data: dict) -> io.BytesIO:
        doc = Document()
        title = form_data.get('info', {}).get('title', 'Formulario sin título')
        doc.add_heading(title, 0)
        desc = form_data.get('info', {}).get('description', '')
        if desc:
            doc.add_paragraph(desc)
            doc.add_paragraph()
        letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        for i, item in enumerate(form_data.get('items', []), 1):
            q = item.get('title', 'Pregunta sin título')
            doc.add_paragraph(f"{i}. {q}", style='Heading 2')
            question = item.get('questionItem', {}).get('question', {})
            if 'choiceQuestion' in question:
                for j, opt in enumerate(question['choiceQuestion'].get('options', [])):
                    prefix = letters[j] if j < len(letters) else str(j+1)
                    doc.add_paragraph(f"  {prefix}. {opt.get('value', 'Opción')}")
            elif 'textQuestion' in question:
                para = question['textQuestion'].get('paragraph', False)
                doc.add_paragraph("  [Respuesta de texto largo]" if para else "  [Respuesta de texto corto]")
            doc.add_paragraph()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf