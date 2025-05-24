# ---------------------------------------------------------------
# Script de Exportacion de Archivos de Google Drive
# Autor: Kevin Gonzalez
# Descripcion:
#   Este script permite exportar archivos de Google Docs, Sheets, Slides
#   y Formularios desde una cuenta de Google Drive a formatos compatibles
#   (.docx, .xlsx, .pptx), incluyendo la conversion de formularios a Word.
#   Utiliza la API de Google Drive y Forms junto con autenticacion OAuth 2.0.
# ---------------------------------------------------------------


import os, io, re, pickle
from tqdm import tqdm
from docx import Document
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

# Funcion para obtener credenciales de Google API
def obtener_credenciales():
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as f:
            cred = pickle.load(f)
    else:
        cred = None
    if not cred or not cred.valid:
        if cred and cred.expired and cred.refresh_token:
            cred.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            cred = flow.run_local_server(port=8089)
        with open('token.pickle', 'wb') as f:
            pickle.dump(cred, f)
    return cred

# Funcion para nombre del archivo
def sanitizar(nombre):
    return re.sub(r'[\\/*?:"<>|]', "_", nombre)

# Funcion genereal para exportar archivos de Google Drive Docs - Sheets - Slides
def exportar_archivos(mime, export_mime, carpeta, ext, tipo):
    service = build('drive', 'v3', credentials=obtener_credenciales())
    os.makedirs(carpeta, exist_ok=True)
    page_token, archivos = None, []

    print(f"Buscando {tipo}s...")
    while True:
        r = service.files().list(q=f"mimeType='{mime}'", fields="nextPageToken, files(id, name)",
                                 pageSize=100, pageToken=page_token).execute()
        archivos += r.get('files', [])
        page_token = r.get('nextPageToken')
        if not page_token: break

    if not archivos:
        print(f"No hay {tipo}s.")
        return

    print(f"Encontrados: {len(archivos)} {tipo}s")
    errores = []
    for a in tqdm(archivos, desc=f"Exportando {tipo}s", unit=ext):
        try:
            request = service.files().export_media(fileId=a['id'], mimeType=export_mime)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            while not downloader.next_chunk()[1]: pass
            with open(os.path.join(carpeta, sanitizar(a['name']) + f'.{ext}'), 'wb') as f:
                f.write(fh.getvalue())
        except Exception as e:
            errores.append((a['name'], str(e)))
    if errores:
        print("\nErrores:")
        for nombre, err in errores:
            print(f" - {nombre}: {err}")
    print(f"Exportado a: {carpeta}")

# Crea un archivo Word a partir de un formulario de Google
def crear_word(form, ruta):
    doc = Document()
    doc.add_heading(form.get('info', {}).get('title', 'Sin titulo'), 0)
    letras = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    for i, item in enumerate(form.get('items', []), 1):
        doc.add_paragraph(f"{i}. {item.get('title', 'Sin titulo')}")
        pregunta = item.get('questionItem', {}).get('question', {})
        if 'choiceQuestion' in pregunta:
            for j, op in enumerate(pregunta['choiceQuestion'].get('options', [])):
                doc.add_paragraph(f"{letras[j]}. {op.get('value', 'Opcion')}")
        elif 'textQuestion' in pregunta:
            doc.add_paragraph("\n\n")
    doc.save(ruta)
    tqdm.write(f"Generado: {ruta}")

# Exporta los formularios de Google Forms a Word
def exportar_forms():
    creds = obtener_credenciales()
    drive = build('drive', 'v3', credentials=creds)
    forms = build('forms', 'v1', credentials=creds)
    os.makedirs("gform_mforms", exist_ok=True)
    page_token, formularios = None, []

    print("Buscando formularios...")
    while True:
        r = drive.files().list(q="mimeType='application/vnd.google-apps.form'",
                               fields="nextPageToken, files(id, name)",
                               pageSize=100, pageToken=page_token).execute()
        formularios += [{'id': f['id'], 'name': f['name']} for f in r.get('files', [])]
        page_token = r.get('nextPageToken')
        if not page_token: break

    if not formularios:
        print("No hay formularios.")
        return

    print(f"Formularios encontrados: {len(formularios)}")
    for form in tqdm(formularios, desc="Exportando formularios", unit="form"):
        try:
            detalle = forms.forms().get(formId=form['id']).execute()
            nombre = sanitizar(detalle.get('info', {}).get('title', form['name']))
            crear_word(detalle, f"gform_mforms/{nombre}.docx")
        except Exception as e:
            print(f"Error con '{form['name']}': {e}")


# Ejecucion principal del script
if __name__ == "__main__":
    exportar_archivos('application/vnd.google-apps.spreadsheet',
                      'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                      'sheets_excel', 'xlsx', 'hoja de calculo')
    print("="*50)
    exportar_archivos('application/vnd.google-apps.presentation',
                      'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                      'slides_pptx', 'pptx', 'presentacion')
    print("="*50)
    exportar_archivos('application/vnd.google-apps.document',
                      'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                      'docs_word', 'docx', 'documento')
    print("="*50)
    exportar_forms()
    print("="*50)
    print("Proceso completado.")
