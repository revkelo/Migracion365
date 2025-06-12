# ---------------------------------------------------------------
# Servicio Google Drive y Forms
# Autor: Kevin Gonzalez
# Descripción:
#   Gestiona la autenticación usando credenciales cifradas y permite
#   listar, descargar y exportar archivos de Google Drive y Formularios.
# ---------------------------------------------------------------

import os
import io
import pickle
import logging
import sys
import json
import time
from pathlib import Path
from cryptography.fernet import Fernet
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from docx import Document
from config import GOOGLE_SCOPES, GOOGLE_EXPORT_FORMATS
from utils import limpiar_archivos

KEY = b"HG5GHGW3o9bMUMWUmz7khGjhELzFUJ9W-52s_ZnIC40="

"""
Carga el blob cifrado desde archivo.

- Si el script está congelado con PyInstaller, busca en sys._MEIPASS.
- Lanza FileNotFoundError si no existe.
- Devuelve los bytes cifrados.
"""

def _cargar_encriptado(filename: str = 'credentials.json.enc') -> bytes:

    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_path, filename)
    if not os.path.isfile(path):
        raise FileNotFoundError(f"No encontré el archivo cifrado en {path}")
    with open(path, 'rb') as f:
        return f.read()

"""
Descifra y carga las credenciales JSON.

- Usa Fernet con la clave KEY para descifrar.
- Parsea JSON y retorna el dict.
"""
def _cargar_credenciales(filename: str = 'credentials.json.enc') -> dict:
    blob = _cargar_encriptado(filename)
    f = Fernet(KEY)
    plaintext = f.decrypt(blob)
    return json.loads(plaintext)

"""
Servicio para interactuar con Google Drive y Google Forms.

Funciones principales:
- Autenticación (token almacenado/recuperado desde pickle).
- Listar unidades compartidas y permisos.
- Listar y reconstruir jerarquía de carpetas.
- Descargar archivos nativos de Google (Docs, Sheets, Slides, Forms)
  exportándolos a formatos Office (docx, xlsx, pptx).
- Descargar archivos binarios convencionales.
- Conversión de formularios a Word.
"""
class GoogleService:
    

    """
    Inicializa el servicio:

    - Guarda rutas de credenciales cifradas y token pickle.
    - Inicializa URL (para GUI) como None.
    - Inicializa atributos para Drive, Forms, usuario autenticado.
    - Configura logger con el nombre de la clase.
    - Llama a _setup_services() para autenticar y obtener clientes API.
    """
    def __init__(self,
                 encrypted_credentials: str = 'credentials.json.enc',
                 token_path: str = 'token.pickle'):
        self.encrypted_credentials = encrypted_credentials
        self.token_path = token_path
        self.url = None
        self.drive = None
        self.usuario = None
        self.forms = None
        self.logger = logging.getLogger(self.__class__.__name__)
        self.servicio_setup()




    """
    Configura credenciales y construye los clientes de API de Drive y Forms.

    - Intenta cargar token desde token_path (.pickle).
    - Si existe y es válido, lo usa.
    - Si no existe o está expirado, inicia un flujo OAuth:
      1) Descifrando JSON de credenciales (credentials.json.enc).
      2) Creando un flujo con InstalledAppFlow usando GOOGLE_SCOPES.
      3) Ejecutando run_local_server() para obtener token vía navegador.
    - Guarda el nuevo token en token_path.
    - Luego, construye los clientes: self.drive y self.forms.
    - Intenta obtener el correo del usuario autenticado y lo guarda en self.usuario.
    """
    def servicio_setup(self):
        creds = None

        if os.path.exists(self.token_path):
            try:
                with open(self.token_path, 'rb') as token_file:
                    creds = pickle.load(token_file)
                    self.logger.info("Token cargado desde %s", self.token_path)
            except Exception:
                self.logger.warning("No se pudo cargar el token, se generará uno nuevo")
                creds = None
                
        if not creds or not getattr(creds, 'valid', False):
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
                self.logger.info("Credenciales refrescadas automáticamente")
            else:
               
                config = _cargar_credenciales('credentials.json.enc')
                flow   = InstalledAppFlow.from_client_config(
                    {'installed': config['installed']},
                    scopes=GOOGLE_SCOPES
                )
                creds = flow.run_local_server(port=8089)
                self.logger.info("Autenticación completada con credenciales cifradas")
          
            with open(self.token_path, 'wb') as token_file:
                pickle.dump(creds, token_file)
                self.logger.info("Token guardado en %s", self.token_path)

        self.drive = build('drive', 'v3', credentials=creds)
        self.forms = build('forms', 'v1', credentials=creds)
        self.logger.info("APIs de Drive y Forms listas para usarse")
        
        try:
            about = self.drive.about().get(fields="user").execute()
            self.usuario = about['user']['emailAddress']
            self.logger.info("Usuario autenticado: %s", self.usuario)
        except Exception as e:
            self.usuario = None
            self.logger.error("No se pudo obtener el usuario autenticado: %s", e)



    """
    Traduce roles en inglés a equivalentes en español para permisos de Drive.

    Recibe un rol (string) y retorna su traducción:
      "organizer"     → "Administrador"
      "fileOrganizer" → "Gestor de contenido"
      "writer"        → "Colaborador"
      "commenter"     → "Comentarista"
      "reader"        → "Lector"
    Si no coincide con ninguna clave, retorna "Desconocido".
    """
    def rol_espanol(self, rol: str) -> str:
        equivalencias = {
            "organizer": "Administrador",
            "fileOrganizer": "Gestor de contenido",
            "writer": "Colaborador",
            "commenter": "Comentarista",
            "reader": "Lector"
        }
        return equivalencias.get(rol, "Desconocido")


    """
    Lista todas las Unidades Compartidas (Drives) donde el usuario tiene acceso.

    - Realiza llamadas paginadas a drive.drives().list()
    - Cada respuesta incluye 'drives(id, name)'
    - Acumula en la lista 'unidades' hasta agotar pageToken
    - Retorna la lista de diccionarios con keys 'id' y 'name'
    """
    def listar_unidades_compartidas(self):
        unidades = []
        page_token = None
        while True:
            res = self.drive.drives().list(
                pageSize=100,
                pageToken=page_token,
                fields="nextPageToken, drives(id, name)"
            ).execute()
            unidades.extend(res.get("drives", []))
            page_token = res.get("nextPageToken")
            if not page_token:
                break
        return unidades



    """
    Lista todos los permisos (ACLs) de un archivo o carpeta en Drive.

    - Parámetro 'file_id': el ID del recurso en Drive
    - Se usa supportsAllDrives=True para incluir unidades compartidas
    - Paginado con pageToken; campos: id, type, role, emailAddress, domain
    - Retorna lista de diccionarios con información de permisos
    """
    def listar_permisos(self, file_id: str):
        permisos = []
        page_token = None
        while True:
            res = self.drive.permissions().list(
                fileId=file_id,
                supportsAllDrives=True,
                pageSize=100,
                pageToken=page_token,
                fields="nextPageToken, permissions(id, type, role, emailAddress, domain)"
            ).execute()
            permisos.extend(res.get("permissions", []))
            page_token = res.get("nextPageToken")
            if not page_token:
                break
        return permisos


    """
    Lista todo el contenido (archivos y carpetas) de una Unidad Compartida.

    - Parámetro 'drive_id': el ID de la unidad compartida
    - Realiza una consulta al endpoint files().list() con:
        corpora='drive', driveId=drive_id, includeItemsFromAllDrives=True,
        supportsAllDrives=True, q="trashed = false"
    - Cada objeto en 'files' tiene id, name, mimeType, parents
    - Retorna una lista de diccionarios con esa información
    """
    def listar_contenido_drive(self, drive_id: str):
        archivos = []
        page_token = None
        while True:
            res = self.drive.files().list(
                corpora='drive',
                driveId=drive_id,
                includeItemsFromAllDrives=True,
                supportsAllDrives=True,
                q="trashed = false",
                pageSize=1000,
                pageToken=page_token,
                fields="nextPageToken, files(id, name, mimeType, parents)"
            ).execute()
            archivos.extend(res.get("files", []))
            page_token = res.get("nextPageToken")
            if not page_token:
                break
        return archivos

    """
    Método auxiliar para obtener el correo electrónico del usuario autenticado.

    - Llama a drive.about().get(fields="user") y extrae 'emailAddress'.
    - En caso de error, registra y retorna cadena vacía.
    """
    def obtener_usuario(self) -> str:

        try:
            about = self.drive.about().get(fields="user").execute()
            return about['user']['emailAddress']
        except Exception as e:
            self.logger.error("No se pudo obtener el usuario: %s", e)
            return ""



    """
    Lista recursivamente todos los archivos y carpetas en el Drive personal (no compartido).

    - Omite elementos en la papelera (q="trashed = false").
    - Pide campos: id, name, mimeType, parents, size, modifiedTime.
    - Clasifica en 'folders' (mimeType carpeta) y 'files' (otros mimeType).
    - Suma el tamaño total de archivos (total_size).
    - Retorna tuplas (folders_dict, files_dict, total_size_bytes).
      donde:
        folders_dict[id] = {id, name, mimeType, parents}
        files_dict[id]   = {id, name, mimeType, parents, size, modifiedTime}
    """
    def listar_archivos_y_carpetas(self):

        folders, files = {}, {}
        page_token = None
        total_size = 0
        while True:
            res = self.drive.files().list(
                q="trashed = false",
                fields="nextPageToken, files(id, name, mimeType, parents, size, modifiedTime)",
                pageSize=1000,
                pageToken=page_token
            ).execute()
            for item in res.get('files', []):
                if item['mimeType'] == 'application/vnd.google-apps.folder':
                    folders[item['id']] = item
                else:
                    files[item['id']] = item
                    total_size += int(item.get('size', 0) or 0)
            page_token = res.get('nextPageToken')
            if not page_token:
                break
        return folders, files, total_size

    """
    Reconstruye la ruta completa de carpetas dado un parent_id.

    - Parámetro 'parent_id': ID de la carpeta actual.
    - Parámetro 'folders': diccionario de carpetas (id → metadatos).
    - Recorre recursivamente hacia arriba (usando parents[0]) hasta raíz.
    - Sanitiza cada nombre de carpeta para evitar caracteres inválidos.
    - Devuelve una lista de nombres en orden desde raíz hasta la carpeta dada.
    """
    def obtener_ruta_carpeta(self, parent_id: str, folders: dict) -> list:

        path, current = [], parent_id
        while current in folders:
            folder = folders[current]
            path.append(limpiar_archivos(folder['name']))
            parents = folder.get('parents') or []
            current = parents[0] if parents else None
        return list(reversed(path))
    
    
    
    def listar_compartidos_conmigo(self):
        """Devuelve la lista de archivos y carpetas que se han compartido conmigo (no unidades)."""
        archivos = []
        page_token = None
        while True:
            resp = self.drive.files().list(
                q="sharedWithMe = true and trashed = false",
                pageSize=1000,
                pageToken=page_token,
                fields="nextPageToken, files(id, name, mimeType, parents, size, modifiedTime)"
            ).execute()
            archivos.extend(resp.get("files", []))
            page_token = resp.get("nextPageToken")
            if not page_token:
                break
        return archivos

    
    
    """
    Descarga o exporta un archivo de Google Drive según su MIME type.

    - file_info: dict con keys 'id', 'mimeType', 'name'.
    - Si el MIME pertenece a GOOGLE_EXPORT_FORMATS, realiza export (Docs, Sheets, Slides).
    - Si es Forms, invoca _create_word_from_form() y retorna un BytesIO con el .docx del formulario.
    - Si es otro tipo (imagen, pdf, etc.), realiza get_media() para descargar bytes.
    - Implementa reintentos (max 3) en caso de errores transitorios (timeout, 500, SSL).
    - Si el tamaño del archivo supera 100MB y es exportable, asigna last_error y retorna (None, name).
    - Retorna tupla (BytesIO, filename) si tuvo éxito, o (None, name) en caso de falla.
    """
    def descargar(self, file_info: dict):
        file_id   = file_info['id']
        mime      = file_info['mimeType']
        name      = file_info['name']
        
 
        if mime in GOOGLE_EXPORT_FORMATS:
            try:
                meta = self.drive.files().get(fileId=file_id, fields="size").execute()
                size_bytes = int(meta.get('size', 0) or 0)
    
                if size_bytes > 100 * 1024 * 1024:
                    self.last_error = Exception("exportSizeLimitExceeded")
                    return None, name
            except Exception:
  
                pass

        max_retries = 3
        attempt = 0
        while attempt < max_retries:
            try:
                if mime in GOOGLE_EXPORT_FORMATS:
                    exp = GOOGLE_EXPORT_FORMATS[mime]
                    if mime == 'application/vnd.google-apps.form':
                        form_data = self.forms.forms().get(formId=file_id).execute()
                        return self.crear_form(form_data), f"{limpiar_archivos(name)}_form.docx"
                    req = self.drive.files().export_media(fileId=file_id, mimeType=exp['mime'])
                else:
                    req = self.drive.files().get_media(fileId=file_id)

                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, req)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)

                ext = exp['ext'] if mime in GOOGLE_EXPORT_FORMATS else None
                filename = f"{limpiar_archivos(name)}.{ext}" if ext else limpiar_archivos(name)
                return fh, filename

            except Exception as e:
                raw = str(e).lower()

                if any(keyword in raw for keyword in ("timed out", "timeout", "500", "ssl")):
                    attempt += 1
                    backoff = 2 ** attempt
                    self.logger.warning("Reintentando descarga de '%s' (intento %d/%d) tras error: %s", name, attempt, max_retries, e)
                    time.sleep(backoff)
                    continue

                self.logger.error("Error irreparable al descargar '%s': %s", name, e)
                self.last_error = e
                return None, name

        self.last_error = Exception("Tiempo de espera agotado tras varios intentos.")
        return None, name


        
    """
    Convierte un formulario de Google Forms a un documento Word (.docx).

    - form_data: dict resultante de forms.forms().get(formId=...)
    - Crea un objeto Document() de python-docx en memoria.
    - Añade título (heading level 0) y descripción si existe.
    - Recorre cada ítem del formulario:
        • Si es pregunta de elección (choiceQuestion), lista opciones con prefijo A., B., C., ...
        • Si es texto (textQuestion), añade placeholder de respuesta corta/larga.
    - Devuelve un BytesIO con el contenido del .docx.
    """
    def crear_form(self, form_data: dict) -> io.BytesIO:

        doc = Document()
        info = form_data.get('info', {})
        doc.add_heading(info.get('title', 'Formulario'), level=0)
        if info.get('description'):
            doc.add_paragraph(info['description'])
            doc.add_paragraph()

        letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        for i, item in enumerate(form_data.get('items', []), start=1):
            q = item.get('title', f'Pregunta {i}')
            doc.add_paragraph(f"{i}. {q}", style='Heading 2')
            question = item.get('questionItem', {}).get('question', {})
            if 'choiceQuestion' in question:
                for j, opt in enumerate(question['choiceQuestion'].get('options', [])):
                    prefix = letters[j] if j < len(letters) else str(j+1)
                    doc.add_paragraph(f"    {prefix}. {opt.get('value', 'Opción')}")
            elif 'textQuestion' in question:
                para = question['textQuestion'].get('paragraph', False)
                placeholder = '[Respuesta de texto largo]' if para else '[Respuesta de texto corto]'
                doc.add_paragraph(f"    {placeholder}")
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf